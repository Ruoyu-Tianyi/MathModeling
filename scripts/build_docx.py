#!/usr/bin/env python3
"""Build a CUMCM-standard .docx paper from paper.md, ON TOP OF the official
template assets/cumcm-template.docx (V1.1).

The template contributes: A4 page & margins (T/B 2.54, L/R 2.70 cm), its own
paragraph styles (Heading 1/2/3 黑体, Normal TNR+宋体 12pt, 图表标题 10.5pt
bold), the 三线表 table style, and the centered page-number footer. This
script empties the template body and re-writes the paper with those styles.

Fully offline: math is rendered with matplotlib mathtext to embedded PNGs
(\tag{n} stripped, re-attached as a right-aligned equation number).

Usage:
    python build_docx.py <paper.md> [--out paper.docx] [--template PATH]

Markdown subset: #..#### headings, paragraphs, **bold**, - and 1. lists,
| tables |, ![caption](path), $$..$$ display math, $..$ inline math.
"""
import argparse
import re
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# OMML (native Word equation) pipeline: LaTeX -> MathML (latex2mathml)
# -> OMML (Office-bundled MML2OMML.XSL). Falls back to mathtext PNG.
import latex2mathml.converter
from lxml import etree

MML2OMML = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_omml_tf = None
if Path(MML2OMML).is_file():
    _omml_tf = etree.XSLT(etree.parse(MML2OMML))
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def preprocess_latex(latex):
    tag = None
    m = TAG_RE.search(latex)
    if m:
        tag = m.group(1)
        latex = TAG_RE.sub("", latex)
    latex = latex.replace(r"\dfrac", r"\frac").strip()
    latex = re.sub(r"\\le(?![a-zA-Z])", r"\\leq", latex)
    latex = re.sub(r"\\ge(?![a-zA-Z])", r"\\geq", latex)
    return latex, tag


def latex_to_omml(latex):
    """Returns m:oMath lxml element, or None on failure."""
    if _omml_tf is None:
        return None
    try:
        mml = latex2mathml.converter.convert(latex)
        omml = _omml_tf(etree.fromstring(mml.encode()))
        return omml.getroot()
    except Exception:
        return None

TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "cumcm-template.docx"

# manual-run fonts for elements the template formats manually (title/摘要头)
TITLE_FONT = {"ea": "黑体", "ascii": "Times New Roman", "size": 16}
ABSH_FONT = {"ea": "黑体", "ascii": "Times New Roman", "size": 14}
BODY_SIZE = 12

MATH_IMG = Path(__file__).resolve().parent / "_math_tmp"
MATH_IMG.mkdir(exist_ok=True)
_math_counter = 0


def set_font(run, font, bold=None):
    run.font.name = font["ascii"]
    run.font.size = Pt(font["size"])
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font["ea"])


TAG_RE = re.compile(r"\\tag\{([^}]*)\}")


def math_png(latex, fontsize=12):
    global _math_counter
    latex, tag = preprocess_latex(latex)
    _math_counter += 1
    path = MATH_IMG / f"m{_math_counter}.png"
    fig = plt.figure(figsize=(0.1, 0.1))
    t = fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.canvas.draw()
    bbox = t.get_window_extent()
    plt.close(fig)
    fig = plt.figure(figsize=(max(bbox.width, 2) / 72, max(bbox.height, 2) / 72))
    fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.savefig(path, dpi=300, bbox_inches="tight", pad_inches=0.01, transparent=True)
    plt.close(fig)
    from PIL import Image
    with Image.open(path) as im:
        w, h = im.size
    return str(path), w, h, tag


INLINE_MATH_RE = re.compile(r"\$([^$]+)\$")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
CAPTION_RE = re.compile(r"^(图|表)\s*\d+\s*[:：]")


def add_runs(p, text, base_size=BODY_SIZE):
    """Add runs handling $math$ and **bold**; style/fonts come from paragraph style."""
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            _add_bold_aware(p, text[pos:m.start()])
        latex, _ = preprocess_latex(m.group(1))
        omml = latex_to_omml(latex)
        if omml is not None:  # native Word equation
            p._p.append(omml)
        else:                 # fallback: mathtext PNG
            img, w, h, _ = math_png(m.group(1), fontsize=int(base_size))
            run = p.add_run()
            height_pt = base_size * 1.15
            run.add_picture(img, height=Pt(height_pt), width=Pt(height_pt * w / h))
        pos = m.end()
    if pos < len(text):
        _add_bold_aware(p, text[pos:])


CODE_RE = re.compile(r"`([^`]+)`")


def _add_bold_aware(p, text):
    def emit(seg):
        pos2 = 0
        for cm in CODE_RE.finditer(seg):
            if cm.start() > pos2:
                p.add_run(seg[pos2:cm.start()])
            r = p.add_run(cm.group(1))
            r.font.name = "Consolas"
            pos2 = cm.end()
        if pos2 < len(seg):
            p.add_run(seg[pos2:])
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            emit(text[pos:m.start()])
        r = p.add_run(re.sub(CODE_RE, r"\1", m.group(1)))
        r.bold = True
        pos = m.end()
    if pos < len(text):
        emit(text[pos:])


def set_cell_border(cell, **edges):
    """edges: edge=(on, size_eighth_pt). e.g. top=(True, 12) -> 1.5pt line."""
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        spec = edges.get(edge)
        if spec and spec[0]:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(spec[1]))
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def set_table_cell_margins(t, side_cm=0.2):
    tbl_pr = t._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 40), ("left", int(side_cm * 567)),
                      ("bottom", 40), ("right", int(side_cm * 567))):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def add_table(doc, rows):
    """3-line table per 优秀作品 spec: 1.5pt top/bottom rules, 0.5pt header
    rule; first column narrow+centered, others left-aligned; cells
    vertically centered with 0.2 cm side margins; borders drawn directly
    (table-style tblLook is unreliable)."""
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    n_rows, n_cols = len(rows), len(rows[0])
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    text_w_cm = 21.0 - 2 * 2.7          # A4 minus template L/R margins
    first_w = 3.0 if n_cols > 1 else text_w_cm
    other_w = (text_w_cm - first_w) / max(n_cols - 1, 1)
    widths = [first_w] + [other_w] * (n_cols - 1)
    set_table_cell_margins(t)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.width = Cm(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, cell_text, base_size=10.5)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(10.5)
                if i == 0:
                    r.bold = True
            set_cell_border(cell,
                            top=(i == 0, 12),            # 顶线 1.5pt
                            bottom=(True, 4) if i == 0 else
                                   ((True, 12) if i == n_rows - 1 else None),
                            # 栏目线 0.5pt / 底线 1.5pt，其余无线
                            left=None, right=None, insideH=None, insideV=None)
    return t


def first_line_indent(p, chars=2):
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    p._p.get_or_add_pPr().append(ind)


def clear_body(doc):
    body = doc.element.body
    for el in list(body):
        if el.tag != qn("w:sectPr"):
            body.remove(el)


def build(md_path: Path, out_path: Path, template: Path):
    doc = Document(str(template))
    clear_body(doc)

    base = md_path.parent
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2)
            if level == 1:  # paper title: template formats it manually
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(p, text, base_size=TITLE_FONT["size"])
                for r in p.runs:
                    if r.text:
                        set_font(r, TITLE_FONT)
            elif level == 2 and text.strip() in ("摘要", "摘  要"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("摘  要")
                set_font(r, ABSH_FONT)
            else:
                # template Heading styles auto-number; strip manual prefixes
                text = re.sub(r"^[一二三四五六七八九十]+、\s*", "", text)
                text = re.sub(r"^\d+(\.\d+)*[、.\s]\s*", "", text)
                style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
                p = doc.add_paragraph(style=style)
                add_runs(p, text)
            i += 1; continue
        if line.strip().startswith("$$"):
            buf = line.strip()[2:]
            while not buf.endswith("$$"):
                i += 1
                buf += lines[i].strip()
            latex, tag = preprocess_latex(buf[:-2])
            omml = latex_to_omml(latex)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if omml is not None:  # native display equation wrapped in oMathPara
                opara = etree.SubElement(p._p, f"{{{M_NS}}}oMathPara")
                opara.append(omml)
                if tag:
                    p.add_run("\t\t（" + tag + "）")
                i += 1; continue
            img, w, h, tag = math_png(buf[:-2], fontsize=13)
            run = p.add_run()
            max_w_cm, nat_w_cm = 14.0, w / 300 * 2.54
            scale = min(1.0, max_w_cm / max(nat_w_cm, 0.1))
            run.add_picture(img, width=Cm(nat_w_cm * scale), height=Cm(h / 300 * 2.54 * scale))
            if tag:
                p.add_run("\t\t（" + tag + "）")
            i += 1; continue
        m = IMG_RE.match(line.strip())
        if m:
            img_path = (base / m.group(2)).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Cm(12.5))
            i += 1; continue
        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c or "---") for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        if CAPTION_RE.match(line.strip()):
            try:
                p = doc.add_paragraph(style="图表标题")
            except KeyError:
                p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line.strip(), base_size=10.5)
            i += 1; continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            try:
                p = doc.add_paragraph(style="List Paragraph")
            except KeyError:
                p = doc.add_paragraph()
            add_runs(p, m.group(2) + " " + m.group(3))
            i += 1; continue
        p = doc.add_paragraph()
        first_line_indent(p, 2)
        add_runs(p, line.strip())
        i += 1

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("paper", help="path to paper.md")
    ap.add_argument("--out", default=None)
    ap.add_argument("--template", default=str(TEMPLATE))
    args = ap.parse_args()
    md = Path(args.paper)
    out = Path(args.out) if args.out else md.with_suffix(".docx")
    build(md, out, Path(args.template))
    print(f"OK: {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    sys.exit(main())
