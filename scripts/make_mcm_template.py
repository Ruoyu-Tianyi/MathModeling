#!/usr/bin/env python3
"""Generate assets/mcm-template.docx — the MCM/ICM base template.

Parameters follow COMAP rules (see references/mcm-format.md):
US Letter, 1-inch margins, Times New Roman 12pt body, bold TNR headings,
"Page X" footer. The docx is a style/page carrier: build_docx.py clears the
body and writes the paper with these named styles.

Run:  python scripts/make_mcm_template.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent.parent / "assets" / "mcm-template.docx"


def style_font(st, name, size, bold):
    st.font.name = name
    st.font.size = Pt(size)
    st.font.bold = bold
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rf.set(qn("w:eastAsia"), name)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)  # US Letter
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)

    style_font(doc.styles["Normal"], "Times New Roman", 12, False)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        style_font(doc.styles[name], "Times New Roman", size, True)

    # caption style ("FigureCaption") if present; else reuse Normal 10.5 bold
    try:
        style_font(doc.styles["Caption"], "Times New Roman", 10.5, True)
        cap = "Caption"
    except KeyError:
        cap = None

    # footer: centered "Page X"
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    p._p.append(fld)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT} ({OUT.stat().st_size} bytes), caption style: {cap}")


if __name__ == "__main__":
    main()
